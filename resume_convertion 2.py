from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
import re
import pathlib

# ---------- Helpers ----------
def read_text(path: str) -> str:
    """Read resume text safely with fallback encoding."""
    try:
        with open(path, "r", encoding="utf-8") as f:
            return f.read().strip()
    except UnicodeDecodeError:
        with open(path, "r", encoding="latin-1") as f:
            return f.read().strip()

def set_margins(section, top=0.5, bottom=0.5, left=0.6, right=0.6):
    section.top_margin = Inches(top)
    section.bottom_margin = Inches(bottom)
    section.left_margin = Inches(left)
    section.right_margin = Inches(right)

def style_base(doc: Document):
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Calibri"
    style.element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    font.size = Pt(10.5)

# ---------- Inline Markdown Parser ----------
BOLD_RE = re.compile(r"\*\*(.*?)\*\*")
ITALIC_RE = re.compile(r"\*(.*?)\*")

def add_formatted_text(paragraph, text: str):
    """Parses inline *italic* and **bold** and adds formatted runs into the paragraph."""
    idx = 0
    while idx < len(text):
        b_match = BOLD_RE.search(text, idx)
        i_match = ITALIC_RE.search(text, idx)
        matches = [m for m in [b_match, i_match] if m]
        if not matches:
            paragraph.add_run(text[idx:])
            break
        m = min(matches, key=lambda x: x.start())
        if m.start() > idx:
            paragraph.add_run(text[idx:m.start()])
        content = m.group(1)
        if m.re == BOLD_RE:
            run = paragraph.add_run(content)
            run.bold = True
        else:
            run = paragraph.add_run(content)
            run.italic = True
        idx = m.end()

# ---------- Resume Builders ----------
def add_name_header(doc: Document, name_line: str, contact_line: str):
    p = doc.add_paragraph()
    add_formatted_text(p, name_line)
    for run in p.runs:
        run.font.size = Pt(20)
        run.bold = True
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    if contact_line.strip():
        p2 = doc.add_paragraph()
        add_formatted_text(p2, contact_line)
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in p2.runs:
            run.font.size = Pt(10.5)

def add_section_heading(doc: Document, title: str):
    p = doc.add_paragraph()
    add_formatted_text(p, title.upper())
    for run in p.runs:
        run.bold = True
        run.font.size = Pt(11.5)
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(2)

def add_bullet(doc: Document, text: str):
    p = doc.add_paragraph(style="List Bullet")
    add_formatted_text(p, text)

def add_separator(doc: Document):
    p = doc.add_paragraph()
    p.add_run("")  # blank line
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)

# ---------- Parser ----------
BULLET_RE = re.compile(r"^\s*[-*]\s+(.*)$")

def parse_and_build(doc: Document, text: str):
    lines = [ln.rstrip() for ln in text.splitlines()]

    # Detect name & contact (first two lines before separator)
    idx = 0
    name_line = lines[idx].strip() if idx < len(lines) else "Your Name"
    idx += 1
    contact_line = ""
    if idx < len(lines) and lines[idx].strip() != "---":
        contact_line = lines[idx].strip()
        idx += 1
    add_name_header(doc, name_line, contact_line)

    buf_paragraph = []

    def flush_paragraph():
        if buf_paragraph:
            text_block = " ".join(buf_paragraph).strip()
            if text_block:
                p = doc.add_paragraph()
                add_formatted_text(p, text_block)
                p.paragraph_format.space_after = Pt(2)
            buf_paragraph.clear()

    while idx < len(lines):
        line = lines[idx].strip()
        idx += 1

        if not line:
            flush_paragraph()
            continue
        if line == "---":
            flush_paragraph()
            add_separator(doc)
            continue
        if BULLET_RE.match(line):
            flush_paragraph()
            add_bullet(doc, BULLET_RE.match(line).group(1))
            continue
        # Section headers: lines fully wrapped in ** ... **
        if line.startswith("**") and line.endswith("**") and len(line) > 4:
            flush_paragraph()
            add_section_heading(doc, line.strip(" *"))
            continue
        buf_paragraph.append(line)

    flush_paragraph()

# ---------- Converter ----------
def convert_resume(input_file, output_file):
    doc = Document()
    style_base(doc)
    set_margins(doc.sections[0])

    text = read_text(input_file)
    parse_and_build(doc, text)

    out = pathlib.Path(output_file).expanduser().resolve()
    doc.save(out)
    print(f"[✔] Resume saved to: {out}")


if __name__ == "__main__":
    convert_resume("resume_4293197910.txt", "resume_4293197910.docx")
